import io
import os
import re

import pandas as pd


def load_dataframe(file_path: str, kind: str) -> pd.DataFrame:
    """Read a file into a DataFrame based on its kind/extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".csv" or kind == "csv":
        return _load_csv(file_path)

    if ext in (".xlsx", ".xls") or kind == "excel":
        return pd.read_excel(file_path)

    if ext == ".json" or kind == "json":
        try:
            return pd.read_json(file_path)
        except ValueError:
            return pd.read_json(file_path, lines=True)

    if ext in (".jsonl", ".ndjson") or kind in ("jsonl", "ndjson"):
        return pd.read_json(file_path, lines=True)

    if ext == ".parquet" or kind == "parquet":
        return pd.read_parquet(file_path)

    if ext in (".tsv", ".tab") or kind == "tsv":
        return _load_csv(file_path, sep="\t")

    if ext in (".txt", ".log") or kind in ("text", "txt", "log"):
        return pd.read_csv(file_path, sep=None, engine="python")

    if ext in (".htm", ".html") or kind == "html":
        tables = pd.read_html(file_path)
        if not tables:
            raise ValueError("No tables found in HTML file")
        return tables[0]

    if ext == ".xml" or kind == "xml":
        return pd.read_xml(file_path)

    if ext == ".pdf" or kind == "pdf":
        return _load_pdf(file_path)

    if ext == ".docx" or kind == "docx":
        return _load_docx(file_path)

    if ext == ".doc" or kind == "doc":
        raise ValueError("Legacy .doc format is not supported; please save as .docx")

    raise ValueError(f"Unsupported file type: {ext!r} (kind={kind!r})")


def _load_csv(file_path: str, sep: str = ",") -> pd.DataFrame:
    """Try common encodings so accented-character CSV files don't crash."""
    for enc in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            return pd.read_csv(file_path, sep=sep, encoding=enc)
        except UnicodeDecodeError:
            continue
    return pd.read_csv(file_path, sep=sep, encoding="utf-8", errors="replace")


def _is_real_table(header: list[str], rows: list[list[str]]) -> bool:
    """
    Validate that an extracted table is real structured data, not garbled
    text fragments that pdfplumber misidentified as a table.
    """
    if len(header) < 2 or len(rows) < 1:
        return False

    auto_names = sum(1 for h in header if re.match(r'^col_\d+$', h))
    short_names = sum(1 for h in header if len(str(h).strip()) <= 2)
    long_names = sum(1 for h in header if len(str(h).strip()) > 30)
    bad_count = auto_names + short_names + long_names

    # If more than 1/3 of headers are garbage, reject
    if bad_count / len(header) > 0.35:
        return False

    # Check that rows have meaningful data in most columns
    filled_ratios = []
    for row in rows[:20]:
        filled = sum(1 for cell in row if str(cell).strip())
        filled_ratios.append(filled / len(row) if len(row) > 0 else 0)
    if filled_ratios:
        avg_filled = sum(filled_ratios) / len(filled_ratios)
        if avg_filled < 0.3:
            return False

    return True


def _load_pdf(file_path: str) -> pd.DataFrame:
    """
    Extract structured data from a PDF.

    Strategy:
    1. Collect every table pdfplumber finds with default settings.
    2. Validate that extracted tables are real data (not garbled text).
    3. If no valid tables, fall back to page text for the text engine.
    """
    import pdfplumber

    table_frames: list[pd.DataFrame] = []
    page_texts: list[str] = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            raw_tables = page.extract_tables()
            for tbl in (raw_tables or []):
                if not tbl or len(tbl) < 2:
                    continue
                header = [str(c).strip() if c else f"col_{i}" for i, c in enumerate(tbl[0])]
                rows = [[str(cell).strip() if cell is not None else "" for cell in row] for row in tbl[1:]]

                if not _is_real_table(header, rows):
                    continue

                df = pd.DataFrame(rows, columns=header)
                df = df.loc[:, df.replace("", None).notna().any()]
                if not df.empty:
                    table_frames.append(df)

            text = page.extract_text()
            if text:
                page_texts.append(text.strip())

    if table_frames:
        combined = pd.concat(table_frames, ignore_index=True)
        _coerce_numeric(combined)
        return combined

    # No valid tables found — return page text for the text engine
    if page_texts:
        return pd.DataFrame({"page": range(1, len(page_texts) + 1), "content": page_texts})

    raise ValueError("Could not extract any content from the PDF")


def _coerce_numeric(df: pd.DataFrame) -> None:
    """In-place coerce string columns to numeric where possible."""
    for col in df.columns:
        try:
            df[col] = pd.to_numeric(df[col])
        except (ValueError, TypeError):
            pass


def _load_docx(file_path: str) -> pd.DataFrame:
    """
    Extract structured data from a Word (.docx) file.

    Strategy:
    1. If the document contains tables, return all tables concatenated.
    2. Otherwise return every non-empty paragraph as a row.
    """
    from docx import Document

    doc = Document(file_path)

    if doc.tables:
        frames = []
        for table in doc.tables:
            if not table.rows:
                continue
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows[1:]
            ]
            if _is_real_table(headers, rows):
                df = pd.DataFrame(rows, columns=headers)
                if not df.empty:
                    frames.append(df)
        if frames:
            combined = pd.concat(frames, ignore_index=True)
            _coerce_numeric(combined)
            return combined

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        raise ValueError("No readable content found in the Word document")
    return pd.DataFrame({"content": paragraphs})


def dataframe_preview(df: pd.DataFrame, rows: int = 50) -> list[dict]:
    """Return a JSON-safe preview of the first `rows` rows."""
    head = df.head(rows)
    safe = head.where(pd.notnull(head), None)
    return safe.to_dict(orient="records")
